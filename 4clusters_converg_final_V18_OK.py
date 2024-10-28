#!/usr/bin/env python
# coding: utf-8
"""
Programa: 4clusters_converg_final_V18_OK.py
Simulação de uma Rede de Sensores Inteligentes em clusters, baseada no método K-Means e na Lei de Friis, com controle de eficiência energética.
@author: Maurício Brigato
Mestrado em Engenharia Elétrica FESJ ICTS
Unesp Universidade Estadual Paulista "Júlio de Mesquita Filho"
https://www2.unesp.br/
https://www.sorocaba.unesp.br/#!/pos-graduacao/--engenharia-eletrica-local/
date: 25/10/2024
"""
import random
import math
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from datetime import datetime
from docx import Document
import os

# Funções de Cálculo e Geração
def minkowski_distance(point, center, lamb):
    return sum(abs(point[k] - center[k]) ** lamb for k in range(len(point))) ** (1 / lamb)

def generate_points(tot_sensors, limit, seed=42):
    random.seed(seed)
    return [(random.uniform(0, limit), random.uniform(0, limit)) for _ in range(tot_sensors)]

def calculate_distance(x1, x2, y1, y2):
    return ((x1 - x2) ** 2 + (y1 - y2) ** 2) ** 0.5

def friis_transmission_power(pt, gt, gr, lambda_, d):
    if d == 0:
        return float('inf')
    return pt * (gt * gr * (lambda_ / (4 * math.pi * d)) ** 2)

def find_cluster_head(cluster, centroid):
    return min(cluster, key=lambda point: calculate_distance(point[0], centroid[0], point[1], centroid[1]))

def save_table_to_word(df, save_path):
    doc = Document()
    doc.add_heading("Tabela Comparativa de Clusters", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    doc_path = os.path.join(save_path, 'Tabela_Comparativa_Clusters.docx')
    doc.save(doc_path)
    print(f"Tabela salva em: {doc_path}")

def main():
    inicio = datetime.now().strftime("%d/%m/%Y, %H:%M:%S")
    print("Início:", inicio)

    # Parâmetros da Simulação
    FREQ = 2.4 * 10**9
    PrdBm = -80
    Pt_transceiver = 0.001
    Pr = 10 ** (PrdBm / 10) / 1000
    LAMBDA = 299792458 / FREQ
    GT = 1.0
    GR = 1.0
    area_limit = 100
    save_path = r'C:\Users\mauri\Documents\Mestrado_OFICIAL\DEFESA MESTRADO\4clustersfinal'
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    # Configuração da Base Station
    sink_position = (100, 50)
    totSensors = 128
    points = generate_points(totSensors, area_limit)

    # Inicialização dos Centroides
    centroids = [(random.uniform(0, area_limit), random.uniform(0, area_limit)) for _ in range(4)]
    trajectory = [[] for _ in range(4)]
    iteration = 0
    tolerance = 1e-3

    # Algoritmo K-Means até Convergência
    while True:
        iteration += 1
        clusters = [[] for _ in range(4)]
        for x, y in points:
            distances = [minkowski_distance((x, y), centroid, 1) for centroid in centroids]
            cluster_idx = distances.index(min(distances))
            clusters[cluster_idx].append((x, y))
        new_centroids = []
        for i, cluster in enumerate(clusters):
            if cluster:
                mean_x = sum(x for x, y in cluster) / len(cluster)
                mean_y = sum(y for x, y in cluster) / len(cluster)
                new_centroids.append((mean_x, mean_y))
                trajectory[i].append((mean_x, mean_y))
            else:
                new_centroids.append((random.uniform(0, area_limit), random.uniform(0, area_limit)))
                trajectory[i].append(centroids[i])
        if all(math.isclose(new_centroids[i][0], centroids[i][0], abs_tol=tolerance) and
               math.isclose(new_centroids[i][1], centroids[i][1], abs_tol=tolerance) for i in range(4)):
            break
        centroids = new_centroids

    # Cálculo das Potências e Eficiência Energética
    cluster_powers = []
    for i, cluster in enumerate(clusters):
        if cluster:
            centroid = centroids[i]
            cluster_head = find_cluster_head(cluster, centroid)
            total_transmission_power = sum(friis_transmission_power(Pt_transceiver, GT, GR, LAMBDA, calculate_distance(sensor[0], cluster_head[0], sensor[1], cluster_head[1])) for sensor in cluster if calculate_distance(sensor[0], cluster_head[0], sensor[1], cluster_head[1]) > 0)
            d_to_sink = calculate_distance(cluster_head[0], sink_position[0], cluster_head[1], sink_position[1])
            ch_to_sink_power = friis_transmission_power(Pt_transceiver, GT, GR, LAMBDA, d_to_sink)
            efficiency = ch_to_sink_power / total_transmission_power
            cluster_powers.append({
                'cluster_index': i + 1,
                'cluster_head': cluster_head,
                'total_transmission_power': total_transmission_power,
                'ch_to_sink_power': ch_to_sink_power,
                'efficiency': efficiency,
                'distance_to_sink': d_to_sink,
                'total_sensors': len(cluster)
            })

    # Criação da Tabela de Comparação e Salvar em Word
    df_clusters = pd.DataFrame({
        "Cluster": [data['cluster_index'] for data in cluster_powers],
        "Total Sensores": [data['total_sensors'] for data in cluster_powers],
        "Distância do CH ao Sink (m)": [data['distance_to_sink'] for data in cluster_powers],
        "Potência Total Gasta (W)": [data['total_transmission_power'] for data in cluster_powers],
        "Eficiência Energética (CH-Sink / Potência Total)": [data['efficiency'] for data in cluster_powers]
    }).sort_values("Eficiência Energética (CH-Sink / Potência Total)", ascending=True).reset_index(drop=True)
    save_table_to_word(df_clusters, save_path)

    # Gráficos
    # Gráfico de Trajetória e Distribuição dos Clusters
    fig, ax = plt.subplots(figsize=(10, 10))
    colors = ['blue', 'green', 'orange', 'purple']
    markers = ['o', 's', 'D', '^']
    for i, (cluster, traj) in enumerate(zip(clusters, trajectory)):
        x_vals = [x for x, y in cluster]
        y_vals = [y for x, y in cluster]
        ax.scatter(x_vals, y_vals, color=colors[i], marker=markers[i], label=f'Cluster {i+1}')
        traj_x = [x for x, y in traj]
        traj_y = [y for x, y in traj]
        ax.plot(traj_x, traj_y, '--', color=colors[i])
        ax.scatter(traj_x[-1], traj_y[-1], color='black', marker='x', s=100)  # Marca final do centroide
        cluster_head = find_cluster_head(cluster, (traj_x[-1], traj_y[-1]))
        ax.scatter(cluster_head[0], cluster_head[1], color='cyan', edgecolor='black', marker='*', s=150)
        ax.text(cluster_head[0], cluster_head[1], f'CH{i+1}', fontsize=12, verticalalignment='bottom')
    ax.scatter(sink_position[0], sink_position[1], color='red', marker='s', s=150, label='Base Station (Sink)')
    ax.set_title("Distribuição dos Clusters e Trajetória dos Centroides")
    ax.set_xlabel("Eixo X")
    ax.set_ylabel("Eixo Y")
    ax.legend()
    plt.savefig(os.path.join(save_path, "Distribuicao_Trajetoria_Centroides.png"))
    plt.show()

    # Gráfico de Eficiência Energética
    plt.figure(figsize=(8, 6))
    plt.bar(df_clusters['Cluster'], df_clusters['Eficiência Energética (CH-Sink / Potência Total)'], color='purple')
    plt.xlabel("Cluster")
    plt.ylabel("Eficiência Energética (CH-Sink / Potência Total)")
    plt.title("Eficiência Energética dos Clusters")
    plt.xticks(df_clusters['Cluster'])
    plt.tight_layout()
    plt.savefig(os.path.join(save_path, "Eficiência_Energética_Clusters.png"))
    plt.show()

    # Gráfico de Potência de Transmissão dos Cluster Heads para o Sink
    plt.figure(figsize=(8, 6))
    ch_to_sink_powers = [data['ch_to_sink_power'] for data in cluster_powers]
    plt.bar(df_clusters['Cluster'], ch_to_sink_powers, color='cyan')
    plt.xlabel("Cluster")
    plt.ylabel("Potência CH-Sink (W)")
    plt.title("Potência de Transmissão dos Cluster Heads para o Sink")
    plt.tight_layout()
    plt.savefig(os.path.join(save_path, "Potencia_Transmissao_CH_Sink.png"))
    plt.show()

    print("Final:", datetime.now().strftime("%d/%m/%Y, %H:%M:%S"))

if __name__ == "__main__":
    main()
